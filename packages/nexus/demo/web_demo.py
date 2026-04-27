#!/usr/bin/env python3
"""
Rune Protocol — Digital Twin Web Demo

A lightweight web UI that wraps the real DigitalTwin from rune-nexus.
Passkey (WebAuthn) login → DigitalTwin creation → real SDK chat.

Usage:
    # 1. Install deps
    pip install fastapi uvicorn py_webauthn

    # 2. Set your API key in .env or environment
    export GEMINI_API_KEY=AIza...

    # 3. Run the demo
    cd rune-nexus
    python demo/web_demo.py

    # 4. Open browser (MUST be localhost or HTTPS for WebAuthn)
    http://localhost:8000

Options (all configurable via .env):
    --port 8000               Server port
    --provider gemini         LLM provider (gemini/openai/anthropic)
    --private-key 0x...       Enable chain mode (BSC + Greenfield)
"""

from __future__ import annotations

import argparse
import asyncio
import json
import logging
import os
import sys
import warnings

# Suppress third-party deprecation warnings
warnings.filterwarnings("ignore", message="authlib.jose module is deprecated")
warnings.filterwarnings("ignore", message=".*EXPERIMENTAL.*PLUGGABLE_AUTH.*")

logger = logging.getLogger(__name__)
import time
import uuid
from pathlib import Path

# ── Make sure nexus is importable ──
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

# ── Load .env ──
def _load_dotenv():
    for d in [PROJECT_ROOT, Path.cwd()]:
        env_path = d / ".env"
        if env_path.is_file():
            with open(env_path) as f:
                for line in f:
                    line = line.strip()
                    if not line or line.startswith("#") or "=" not in line:
                        continue
                    key, _, value = line.partition("=")
                    key, value = key.strip(), value.strip()
                    if key and key not in os.environ:
                        os.environ[key] = value
            return str(env_path)
    return None

_load_dotenv()

# ── Now import FastAPI and nexus ──
try:
    from fastapi import FastAPI
    from fastapi.middleware.cors import CORSMiddleware
    from fastapi.responses import HTMLResponse, JSONResponse
    from pydantic import BaseModel
    import uvicorn
except ImportError:
    print("Missing dependency: pip install fastapi uvicorn")
    sys.exit(1)

from nexus import DigitalTwin

# ── WebAuthn (passkey) ──
try:
    from webauthn import (
        generate_registration_options,
        verify_registration_response,
        generate_authentication_options,
        verify_authentication_response,
        options_to_json,
    )
    from webauthn.helpers.structs import (
        AuthenticatorSelectionCriteria,
        ResidentKeyRequirement,
        UserVerificationRequirement,
        PublicKeyCredentialDescriptor,
    )
    from webauthn.helpers import bytes_to_base64url, parse_registration_credential_json, parse_authentication_credential_json
    HAS_WEBAUTHN = True
except ImportError:
    HAS_WEBAUTHN = False
    print("⚠️  py_webauthn not installed. Run: pip install py_webauthn")
    print("   Passkey login will be disabled.\n")


# ═══════════════════════════════════════════════════
# Global state
# ═══════════════════════════════════════════════════
twins: dict[str, DigitalTwin] = {}      # session_id -> twin
events: dict[str, list] = {}            # session_id -> on-chain events
twin_lock = asyncio.Lock()
challenge_lock = asyncio.Lock()

# ── User store (in-memory, demo only) ──
# user_id -> { display_name, credentials: [...], agent_id, session_id (if active) }
users: dict[str, dict] = {}
# Pending challenges for registration/authentication
challenges: dict[str, bytes] = {}       # user_id or "auth" -> challenge bytes

# WebAuthn config — updated in main() based on actual port
RP_ID = os.environ.get("WEBAUTHN_RP_ID", "localhost")
RP_NAME = "Rune Protocol"
ORIGIN = os.environ.get("WEBAUTHN_ORIGIN", "http://localhost:8000")

# Config (resolved at startup)
CFG = {}

# ── User persistence ──
USERS_FILE = PROJECT_ROOT / ".nexus_demo" / "users.json"


def _save_users():
    """Persist users to disk (credentials stored as hex for JSON compat)."""
    import base64
    USERS_FILE.parent.mkdir(parents=True, exist_ok=True)
    serializable = {}
    for uid, user in users.items():
        creds = []
        for c in user.get("credentials", []):
            creds.append({
                "id": base64.b64encode(c["id"]).decode() if isinstance(c["id"], bytes) else c["id"],
                "public_key": base64.b64encode(c["public_key"]).decode() if isinstance(c["public_key"], bytes) else c["public_key"],
                "sign_count": c["sign_count"],
            })
        serializable[uid] = {
            "display_name": user["display_name"],
            "agent_id": user["agent_id"],
            "credentials": creds,
        }
    USERS_FILE.write_text(json.dumps(serializable, indent=2))


def _load_users():
    """Load users from disk on startup."""
    import base64
    if not USERS_FILE.exists():
        return
    try:
        data = json.loads(USERS_FILE.read_text())
        for uid, user in data.items():
            creds = []
            for c in user.get("credentials", []):
                creds.append({
                    "id": base64.b64decode(c["id"]) if isinstance(c["id"], str) else c["id"],
                    "public_key": base64.b64decode(c["public_key"]) if isinstance(c["public_key"], str) else c["public_key"],
                    "sign_count": c.get("sign_count", 0),
                })
            users[uid] = {
                "display_name": user["display_name"],
                "agent_id": user["agent_id"],
                "credentials": creds,
            }
        logger.info("Loaded %d users from %s", len(users), USERS_FILE)
    except Exception as e:
        logger.warning("Failed to load users: %s", e)


# ═══════════════════════════════════════════════════
# Event handler — captures on-chain activity
# ═══════════════════════════════════════════════════
def make_event_handler(session_id: str):
    """Create an on_event callback that stores events for a session."""
    def handler(event_type: str, detail: dict):
        if session_id not in events:
            events[session_id] = []
        events[session_id].append({
            "type": event_type,
            "detail": detail,
            "timestamp": time.time(),
        })
    return handler


# ═══════════════════════════════════════════════════
# FastAPI App
# ═══════════════════════════════════════════════════
app = FastAPI(title="Rune Protocol — Digital Twin Demo")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# File storage directories
UPLOAD_DIR = PROJECT_ROOT / ".nexus_demo" / "uploads"
OUTPUT_DIR = PROJECT_ROOT / ".nexus_demo" / "outputs"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# Users loaded in main() after config is resolved


class CreateRequest(BaseModel):
    name: str = ""
    owner: str = ""


class ChatRequest(BaseModel):
    session_id: str
    message: str


# ═══════════════════════════════════════════════════
# Passkey (WebAuthn) endpoints
# ═══════════════════════════════════════════════════

class PasskeyRegisterStartRequest(BaseModel):
    display_name: str


class PasskeyRegisterFinishRequest(BaseModel):
    user_id: str
    credential: str  # JSON string from @simplewebauthn/browser


class PasskeyLoginFinishRequest(BaseModel):
    credential: str  # JSON string from @simplewebauthn/browser


@app.post("/api/passkey/register/start")
async def passkey_register_start(req: PasskeyRegisterStartRequest):
    """Generate registration options for a new passkey."""
    if not HAS_WEBAUTHN:
        return JSONResponse(status_code=501, content={"error": "Passkey not available"})

    user_id = uuid.uuid4().bytes
    user_id_b64 = bytes_to_base64url(user_id)

    options = generate_registration_options(
        rp_id=RP_ID,
        rp_name=RP_NAME,
        user_id=user_id,
        user_name=req.display_name,
        user_display_name=req.display_name,
        authenticator_selection=AuthenticatorSelectionCriteria(
            resident_key=ResidentKeyRequirement.REQUIRED,
            user_verification=UserVerificationRequirement.PREFERRED,
        ),
    )

    async with challenge_lock:
        challenges[user_id_b64] = options.challenge
        users[user_id_b64] = {
            "display_name": req.display_name,
            "credentials": [],
            "agent_id": f"twin-{user_id_b64[:12]}",
        }

    return JSONResponse(content={
        "user_id": user_id_b64,
        "options": json.loads(options_to_json(options)),
    })


@app.post("/api/passkey/register/finish")
async def passkey_register_finish(req: PasskeyRegisterFinishRequest):
    """Verify registration and store credential."""
    if not HAS_WEBAUTHN:
        return JSONResponse(status_code=501, content={"error": "Passkey not available"})

    user_id = req.user_id

    async with challenge_lock:
        challenge = challenges.pop(user_id, None)
    if not challenge:
        return JSONResponse(status_code=400, content={"error": "No pending challenge"})

    try:
        credential = parse_registration_credential_json(req.credential)
        verification = verify_registration_response(
            credential=credential,
            expected_challenge=challenge,
            expected_rp_id=RP_ID,
            expected_origin=ORIGIN,
        )
    except Exception as e:
        logger.warning("Passkey registration failed: %s", e)
        return JSONResponse(status_code=400, content={"error": "Registration verification failed"})

    user = users.get(user_id)
    if not user:
        return JSONResponse(status_code=400, content={"error": "User not found"})

    user["credentials"].append({
        "id": verification.credential_id,
        "public_key": verification.credential_public_key,
        "sign_count": verification.sign_count,
    })
    _save_users()

    return {
        "status": "ok",
        "user_id": user_id,
        "agent_id": user["agent_id"],
        "display_name": users[user_id]["display_name"],
    }


@app.post("/api/passkey/login/start")
async def passkey_login_start():
    """Generate authentication options for passkey login."""
    if not HAS_WEBAUTHN:
        return JSONResponse(status_code=501, content={"error": "Passkey not available"})

    options = generate_authentication_options(
        rp_id=RP_ID,
        user_verification=UserVerificationRequirement.PREFERRED,
    )

    async with challenge_lock:
        challenges["_login_"] = options.challenge

    return JSONResponse(content={
        "options": json.loads(options_to_json(options)),
    })


@app.post("/api/passkey/login/finish")
async def passkey_login_finish(req: PasskeyLoginFinishRequest):
    """Verify authentication and return user info."""
    if not HAS_WEBAUTHN:
        return JSONResponse(status_code=501, content={"error": "Passkey not available"})

    async with challenge_lock:
        challenge = challenges.pop("_login_", None)
    if not challenge:
        return JSONResponse(status_code=400, content={"error": "No pending challenge"})

    try:
        credential = parse_authentication_credential_json(req.credential)
    except Exception as e:
        logger.warning("Passkey login parse failed: %s", e)
        return JSONResponse(status_code=400, content={"error": "Invalid credential format"})

    # Find user with matching credential
    matched_user_id = None
    matched_cred = None
    for uid, user in users.items():
        for cred in user.get("credentials", []):
            if cred["id"] == credential.raw_id:
                matched_user_id = uid
                matched_cred = cred
                break
        if matched_user_id:
            break

    if not matched_user_id or not matched_cred:
        return JSONResponse(status_code=400, content={"error": "Credential not recognized"})

    try:
        verification = verify_authentication_response(
            credential=credential,
            expected_challenge=challenge,
            expected_rp_id=RP_ID,
            expected_origin=ORIGIN,
            credential_public_key=matched_cred["public_key"],
            credential_current_sign_count=matched_cred["sign_count"],
        )
        matched_cred["sign_count"] = verification.new_sign_count
    except Exception as e:
        logger.warning("Passkey auth failed: %s", e)
        return JSONResponse(status_code=400, content={"error": "Authentication failed"})

    user = users[matched_user_id]
    return {
        "status": "ok",
        "user_id": matched_user_id,
        "agent_id": user["agent_id"],
        "display_name": user["display_name"],
    }


# ═══════════════════════════════════════════════════
# User status endpoint
# ═══════════════════════════════════════════════════

@app.get("/api/user/{user_id}")
async def get_user_status(user_id: str):
    """Check if user has an existing twin session."""
    user = users.get(user_id)
    if not user:
        return JSONResponse(status_code=404, content={"error": "User not found"})

    session_id = user.get("session_id")
    has_twin = session_id is not None and session_id in twins

    result = {
        "user_id": user_id,
        "display_name": user["display_name"],
        "agent_id": user["agent_id"],
        "has_twin": has_twin,
    }

    if has_twin:
        twin = twins[session_id]
        result["session_id"] = session_id
        result["message_count"] = len(twin._messages)

    return result


@app.post("/api/logout")
async def logout(req: dict):
    """Logout: close twin but keep user credentials for re-login."""
    user_id = req.get("user_id", "")
    user = users.get(user_id)
    if not user:
        return {"status": "ok"}

    session_id = user.get("session_id")
    if session_id:
        twin = twins.pop(session_id, None)
        if twin:
            await twin.close()
        events.pop(session_id, None)
        user["session_id"] = None

    return {"status": "ok"}


# ═══════════════════════════════════════════════════
# DigitalTwin endpoints
# ═══════════════════════════════════════════════════

@app.post("/api/create")
async def create_twin(req: CreateRequest):
    """Create or resume a DigitalTwin instance with real SDK."""
    # Determine stable agent_id from user mapping
    user = users.get(req.owner) if req.owner else None
    stable_agent_id = user["agent_id"] if user else f"twin-{str(uuid.uuid4())[:8]}"

    # If user already has an active twin, return it
    existing_session = user.get("session_id") if user else None
    if existing_session and existing_session in twins:
        twin = twins[existing_session]
        mem_count = twin.event_log.count()
        recent_mems = []
        try:
            for ev in twin.event_log.recent(limit=10):
                cat = "fact" if ev.event_type == "user_message" else "preference"
                recent_mems.append({"content": ev.content[:100], "category": cat})
        except Exception:
            pass
        return {
            "session_id": existing_session,
            "agent_id": stable_agent_id,
            "chain_info": {"mode": "chain" if CFG.get("private_key") else "local", "network": CFG.get("network", "testnet")},
            "tools": twin.tools.tool_names if twin.tools else [],
            "events": events.get(existing_session, []),
            "resumed": True,
            "message_count": len(twin._messages),
            "memory_count": mem_count,
            "memories": recent_mems,
        }

    session_id = str(uuid.uuid4())[:12]
    events[session_id] = []

    try:
        twin = await DigitalTwin.create(
            name=req.name or CFG.get("name", "Digital Twin"),
            owner=req.owner or CFG.get("owner", "demo-user"),
            agent_id=stable_agent_id,
            llm_provider=CFG["provider"],
            llm_api_key=CFG["api_key"],
            llm_model=CFG.get("model", ""),
            base_dir=str(PROJECT_ROOT / ".nexus_demo" / stable_agent_id),
            # Tools
            enable_tools=True,  # Always enable — FileGenerator doesn't need API keys
            tavily_api_key=CFG.get("tavily_key", ""),
            jina_api_key=CFG.get("jina_key", ""),
            # Chain mode
            private_key=CFG.get("private_key", ""),
            network=CFG.get("network", "testnet"),
            rpc_url=CFG.get("rpc_url", ""),
            greenfield_bucket=CFG.get("greenfield_bucket", "nexus-agent-state"),
            identity_registry_address=CFG.get("identity_registry", ""),
            agent_state_address=CFG.get("agent_state_address", ""),
            task_manager_address=CFG.get("task_manager_address", ""),
        )
    except Exception as e:
        logger.exception("Operation failed")
        return JSONResponse(status_code=500, content={"error": "Internal error. Check server logs."})

    # Wire up event handler
    twin.on_event = make_event_handler(session_id)

    twins[session_id] = twin

    # Link user to this session (if owner matches a user_id)
    if req.owner and req.owner in users:
        users[req.owner]["session_id"] = session_id

    # Collect chain info
    chain_info = {}
    erc_id = getattr(twin, "_erc8004_agent_id", None)
    if erc_id is not None:
        chain_info["erc8004_id"] = str(erc_id)
    chain_info["mode"] = "chain" if CFG.get("private_key") else "local"
    chain_info["network"] = CFG.get("network", "testnet")

    # Collect tool info
    tool_names = twin.tools.tool_names if twin.tools else []

    resumed = bool(twin._messages)
    msg_count = len(twin._messages)
    mem_count = 0
    recent_memories = []
    try:
        mem_count = twin.event_log.count()
        for ev in twin.event_log.recent(limit=10):
            cat = "fact" if ev.event_type == "user_message" else "preference"
            recent_memories.append({"content": ev.content[:100], "category": cat})
    except Exception as e:
        logger.debug("Memory stats error: %s", e)

    logger.info("Twin created: agent=%s, resumed=%s, messages=%d, memories=%d",
                stable_agent_id, resumed, msg_count, mem_count)

    return {
        "session_id": session_id,
        "agent_id": stable_agent_id,
        "chain_info": chain_info,
        "tools": tool_names,
        "events": events.get(session_id, []),
        "resumed": resumed or mem_count > 0,
        "message_count": msg_count,
        "memory_count": mem_count,
        "memories": recent_memories,
    }


@app.post("/api/chat")
async def chat(req: ChatRequest):
    """Send a message to the DigitalTwin and get a response."""
    twin = twins.get(req.session_id)
    if not twin:
        return JSONResponse(status_code=404, content={"error": "Session not found. Create a twin first."})

    # Clear events before this turn
    pre_event_count = len(events.get(req.session_id, []))

    try:
        start = time.time()
        response = await twin.chat(req.message)
        elapsed = time.time() - start
    except Exception as e:
        logging.exception("Chat error")
        logger.exception("Operation failed")
        return JSONResponse(status_code=500, content={"error": "Internal error. Check server logs."})

    # Wait a moment for background tasks (memory extraction, evolution)
    await asyncio.sleep(0.5)

    # Collect new events that fired during this turn
    all_events = events.get(req.session_id, [])
    new_events = all_events[pre_event_count:]

    # Get current stats from the real provider
    memory_count = 0
    skill_count = 0
    recent_memories = []
    try:
        memory_count = twin.event_log.count()
        for ev in twin.event_log.recent(limit=10):
            cat = "fact" if ev.event_type == "user_message" else "preference"
            recent_memories.append({"content": ev.content[:100], "category": cat})
    except Exception as e:
        logger.debug("Memory stats error: %s", e)

    try:
        # Show registered tools count (real capabilities, not learned text skills)
        skill_count = len(twin.tools) if twin.tools else 0
    except Exception:
        pass

    # Check if user's message implies a skill we can suggest
    skill_suggestion = detect_skill_suggestion(
        req.message, twin.skills.names
    )

    return {
        "reply": response,
        "elapsed": round(elapsed, 2),
        "events": [{"type": e["type"], "detail": e["detail"]} for e in new_events],
        "stats": {
            "memory_count": memory_count,
            "skill_count": skill_count,
            "message_count": len(twin._messages),
            "installed_skills": twin.skills.names,
        },
        "memories": recent_memories,
        "skill_suggestion": skill_suggestion,
    }


# ═══════════════════════════════════════════════════
# Skill installation endpoints
# ═══════════════════════════════════════════════════

# Known skills that can be suggested during conversation
SKILL_CATALOG = {
    "wallet": {
        "name": "binance-agentic-wallet",
        "title": "Binance Agentic Wallet",
        "url": "https://github.com/binance/binance-skills-hub/tree/main/skills/binance-web3/binance-agentic-wallet",
        "keywords": ["wallet", "balance", "send", "transfer", "swap", "token", "BNB", "USDT", "crypto", "pay"],
        "description": "Connect wallet, check balance, send tokens, swap on DEX",
    },
    "market": {
        "name": "crypto-market-rank",
        "title": "Crypto Market Rankings",
        "url": "https://github.com/binance/binance-skills-hub/tree/main/skills/binance-web3/crypto-market-rank",
        "keywords": ["market", "price", "rank", "token", "coin", "cap", "volume", "trending"],
        "description": "Search token info, check prices, market rankings",
    },
    "trading": {
        "name": "derivatives-trading-usds-futures",
        "title": "Derivatives Trading",
        "url": "https://github.com/binance/binance-skills-hub/tree/main/skills/binance/derivatives-trading-usds-futures",
        "keywords": ["trade", "futures", "long", "short", "leverage", "position", "order"],
        "description": "Trade USDS futures, manage positions, set orders",
    },
}


class SkillInstallRequest(BaseModel):
    session_id: str
    skill_key: str         # key from SKILL_CATALOG or LobeHub identifier
    source: str = "catalog" # "catalog" or "lobehub"


class SkillSearchRequest(BaseModel):
    session_id: str
    query: str
    limit: int = 5


@app.post("/api/skills/search")
async def search_skills(req: SkillSearchRequest):
    """Search LobeHub marketplace for skills."""
    twin = twins.get(req.session_id)
    if not twin:
        return JSONResponse(status_code=404, content={"error": "Session not found"})

    results = await twin.skills.search_lobehub(req.query, limit=req.limit)
    return {"results": results, "query": req.query}


@app.post("/api/skills/install")
async def install_skill(req: SkillInstallRequest):
    """Install a skill from catalog or LobeHub."""
    twin = twins.get(req.session_id)
    if not twin:
        return JSONResponse(status_code=404, content={"error": "Session not found"})

    try:
        if req.source == "lobehub":
            # Install from LobeHub marketplace
            skill = await twin.skills.install(f"lobehub:{req.skill_key}")
        else:
            # Install from built-in catalog
            catalog_entry = SKILL_CATALOG.get(req.skill_key)
            if not catalog_entry:
                return JSONResponse(status_code=400, content={"error": f"Unknown skill: {req.skill_key}"})
            skill = await twin.skills.install(catalog_entry["url"])

        return {
            "status": "installed",
            "skill_name": skill.name,
            "skill_title": skill.title,
            "description": skill.description,
        }
    except Exception as e:
        logger.warning("Skill install failed: %s", e)
        return JSONResponse(status_code=500, content={"error": f"Install failed: {str(e)[:100]}"})


@app.get("/api/skills/{session_id}")
async def list_skills(session_id: str):
    """List installed skills for a twin."""
    twin = twins.get(session_id)
    if not twin:
        return JSONResponse(status_code=404, content={"error": "Session not found"})

    return {
        "installed": [
            {"name": s.name, "title": s.title, "description": s.description}
            for s in twin.skills.installed
        ],
        "available": [
            {"key": k, "title": v["title"], "description": v["description"]}
            for k, v in SKILL_CATALOG.items()
            if v["name"] not in twin.skills.names
        ],
    }


def detect_skill_suggestion(message: str, installed_names: list[str]) -> dict | None:
    """Check if a user message implies a skill that's not installed.

    Returns a catalog match, or a LobeHub search suggestion for unmatched requests.
    """
    msg_lower = message.lower()

    # 1. Check built-in catalog
    for key, entry in SKILL_CATALOG.items():
        if entry["name"] in installed_names:
            continue
        if any(kw in msg_lower for kw in entry["keywords"]):
            return {
                "skill_key": key,
                "title": entry["title"],
                "description": entry["description"],
                "source": "catalog",
            }

    # 2. Detect capability requests that imply a skill is needed (CN + EN)
    capability_hints = [
        "pdf", "excel", "ppt", "pptx", "docx", "word",
        "email", "邮件", "deploy", "部署",
        "image", "图片", "图像", "video", "视频",
        "输出", "导出", "export", "generate", "生成",
        "创建", "create", "制作", "convert", "转换",
        "帮我", "帮忙", "help me", "i need", "i want",
        "可以", "能不能", "怎么",
    ]
    if any(hint in msg_lower for hint in capability_hints):
        # Extract a search query from the message
        query = msg_lower[:50].replace("帮我", "").replace("帮忙", "").strip()
        return {
            "skill_key": query,
            "title": "Search LobeHub Skills",
            "description": f"Search 100K+ skills for: {message[:60]}",
            "source": "lobehub_search",
        }

    return None


@app.post("/api/close")
async def close_twin(req: dict):
    """Gracefully close a twin session."""
    session_id = req.get("session_id", "")
    twin = twins.pop(session_id, None)
    if twin:
        await twin.close()
        events.pop(session_id, None)
        return {"status": "closed"}
    return {"status": "not_found"}


# ═══════════════════════════════════════════════════
# File upload & download
# ═══════════════════════════════════════════════════

from fastapi import UploadFile, File, Form
from fastapi.responses import FileResponse as FastAPIFileResponse
import shutil


@app.post("/api/upload")
async def upload_file(
    session_id: str = Form(...),
    file: UploadFile = File(...),
):
    """Upload a file and feed its content to the digital twin.

    Text files: content extracted and appended to event log.
    Binary files: stored and referenced by path.
    """
    twin = twins.get(session_id)
    if not twin:
        return JSONResponse(status_code=404, content={"error": "Session not found"})

    # Save file (stream to disk to handle large files)
    MAX_UPLOAD_SIZE = 200 * 1024 * 1024  # 200MB
    safe_name = file.filename.replace("/", "_").replace("\\", "_")
    file_path = UPLOAD_DIR / f"{session_id}_{safe_name}"
    total_written = 0
    with open(file_path, "wb") as f:
        while chunk := await file.read(8192):
            total_written += len(chunk)
            if total_written > MAX_UPLOAD_SIZE:
                file_path.unlink(missing_ok=True)
                return JSONResponse(
                    status_code=413,
                    content={"error": f"File too large. Maximum: {MAX_UPLOAD_SIZE // (1024*1024)}MB"},
                )
            f.write(chunk)

    file_size = file_path.stat().st_size
    logger.info("File uploaded: %s (%d bytes)", safe_name, file_size)

    # Extract content based on file type.
    # full_text: complete extraction for ReadUploadedFileTool (stored in memory or disk).
    # content_preview: truncated (5KB) for event log persistence.
    MAX_EVENT_LOG = 5000     # 5KB for event log
    full_text = ""
    content_preview = ""
    ext = safe_name.rsplit(".", 1)[-1].lower() if "." in safe_name else ""

    if ext in ("txt", "md", "csv", "json", "py", "js", "yaml", "yml", "xml", "html"):
        try:
            text = file_path.read_text(encoding="utf-8", errors="replace")
            full_text = text  # Full content — ReadUploadedFileTool handles storage
            content_preview = text[:MAX_EVENT_LOG]
            twin.event_log.append(
                "file_upload",
                f"[File: {safe_name}] ({file_size} bytes)\n{content_preview}",
                session_id=twin._thread_id,
                metadata={"filename": safe_name, "size": file_size, "type": ext},
            )
        except Exception as e:
            content_preview = f"(text extraction failed: {e})"
    elif ext == "pdf":
        pdf_text = ""
        # Try 1: pdftotext (poppler-utils) — fast, best quality
        try:
            import subprocess
            result = subprocess.run(
                ["pdftotext", str(file_path), "-"],
                capture_output=True, text=True, timeout=120,
            )
            if result.stdout and result.stdout.strip():
                pdf_text = result.stdout
        except Exception:
            pass
        # Try 2: pypdf — pure Python fallback
        if not pdf_text:
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(file_path))
                pages = []
                for i, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(f"[Page {i+1}]\n{text}")
                pdf_text = "\n\n".join(pages)
            except Exception as e:
                logger.warning("pypdf extraction failed: %s", e)
        # Try 3: pdfplumber — handles complex layouts
        if not pdf_text:
            try:
                import pdfplumber
                with pdfplumber.open(str(file_path)) as pdf:
                    pages = []
                    for i, page in enumerate(pdf.pages):
                        text = page.extract_text() or ""
                        if text.strip():
                            pages.append(f"[Page {i+1}]\n{text}")
                    pdf_text = "\n\n".join(pages)
            except Exception as e:
                logger.warning("pdfplumber extraction failed: %s", e)

        if pdf_text:
            full_text = pdf_text
            content_preview = pdf_text[:MAX_EVENT_LOG]
            twin.event_log.append(
                "file_upload",
                f"[PDF: {safe_name}] ({file_size} bytes)\n{content_preview}",
                session_id=twin._thread_id,
                metadata={"filename": safe_name, "size": file_size, "type": "pdf"},
            )
        else:
            content_preview = "(PDF text extraction failed — no extraction library available. Install: pip install pypdf)"
    elif ext in ("docx", "doc"):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(file_path))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    text_parts.append(" | ".join(cells))
            text = "\n".join(text_parts)
            if text:
                full_text = text
                content_preview = text[:MAX_EVENT_LOG]
                twin.event_log.append(
                    "file_upload",
                    f"[DOCX: {safe_name}] ({file_size} bytes)\n{content_preview}",
                    session_id=twin._thread_id,
                    metadata={"filename": safe_name, "size": file_size, "type": ext},
                )
            else:
                content_preview = "(DOCX text extraction returned empty)"
        except ImportError:
            content_preview = "(python-docx not installed — file stored but not extracted)"
        except Exception as e:
            content_preview = f"(DOCX extraction failed: {e})"
    elif ext in ("xlsx", "xls"):
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(file_path), read_only=True, data_only=True)
            rows = []
            for sheet in wb.sheetnames[:5]:  # First 5 sheets
                ws = wb[sheet]
                rows.append(f"[Sheet: {sheet}]")
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i > 200:  # More rows for tool storage
                        rows.append("... (truncated at 200 rows)")
                        break
                    rows.append("\t".join(str(c) if c is not None else "" for c in row))
            wb.close()
            text = "\n".join(rows)
            if text:
                full_text = text
                content_preview = text[:MAX_EVENT_LOG]
                twin.event_log.append(
                    "file_upload",
                    f"[Excel: {safe_name}] ({file_size} bytes)\n{content_preview}",
                    session_id=twin._thread_id,
                    metadata={"filename": safe_name, "size": file_size, "type": ext},
                )
            else:
                content_preview = "(Excel extraction returned empty)"
        except ImportError:
            content_preview = "(openpyxl not installed — file stored but not extracted)"
        except Exception as e:
            content_preview = f"(Excel extraction failed: {e})"
    else:
        # Binary file — just reference it
        twin.event_log.append(
            "file_upload",
            f"[File: {safe_name}] ({file_size} bytes, type: {ext}) — stored at {file_path}",
            session_id=twin._thread_id,
            metadata={"filename": safe_name, "size": file_size, "type": ext, "path": str(file_path)},
        )
        content_preview = f"(binary file, {file_size} bytes)"

    # CRITICAL: Inject file info into _messages so the LLM knows about the file.
    # For large files, store full content in ReadUploadedFileTool — agent reads on demand.
    # For small files (<2KB), inline the full content directly in _messages.
    INLINE_THRESHOLD = 2000  # chars — files smaller than this go fully inline

    is_real_content = full_text and not full_text.startswith("(")

    if is_real_content:
        # Store full text in ReadUploadedFileTool for on-demand reading
        if twin._file_reader:
            twin._file_reader.store(safe_name, full_text)

        if len(full_text) <= INLINE_THRESHOLD:
            # Small file — inline everything
            file_msg = (
                f"[System: The user uploaded '{safe_name}' ({file_size} bytes). "
                f"Full content:]\n\n{full_text}"
            )
        else:
            # Large file — preview + tell agent to use read_uploaded_file tool
            preview = full_text[:500]
            file_msg = (
                f"[System: The user uploaded '{safe_name}' ({file_size} bytes, {len(full_text)} chars extracted). "
                f"Preview (first 500 chars):]\n\n{preview}\n\n"
                f"[The full content is available via the read_uploaded_file tool. "
                f"Call read_uploaded_file(filename='{safe_name}') to read more.]"
            )
        twin._messages.append({"role": "user", "content": file_msg})
        twin._messages.append({
            "role": "assistant",
            "content": f"I've received the file **{safe_name}**. I can see its content and answer questions about it."
        })
    else:
        # Binary or failed extraction — still notify the LLM
        twin._messages.append({"role": "user", "content": f"[System: The user uploaded '{safe_name}' ({file_size} bytes, type: {ext}). Text could not be extracted.]"})
        twin._messages.append({
            "role": "assistant",
            "content": f"I've received the file **{safe_name}**, but I wasn't able to extract its text content."
        })

    # For large files with real content, trigger a background file-digest compact.
    # This produces a curated summary that persists across sessions.
    FILE_DIGEST_THRESHOLD = 3000  # chars — files larger than this get auto-summarized
    if is_real_content and len(full_text) > FILE_DIGEST_THRESHOLD and twin._compactor:
        import asyncio

        async def _file_digest():
            """Background: ask LLM to summarize the file, store in CuratedMemory."""
            try:
                # Use first 15K chars for summarization (fits most LLM context)
                sample = full_text[:15000]
                prompt = (
                    f"The user uploaded a file '{safe_name}' ({len(full_text)} chars). "
                    f"Summarize the key points in 3-5 bullet points:\n\n{sample}"
                )
                summary = await twin.llm.complete(prompt)
                if summary:
                    # Store digest in event log (persists to Greenfield)
                    twin.event_log.append(
                        "file_digest",
                        f"[Digest of {safe_name}]\n{summary}",
                        session_id=twin._thread_id,
                        metadata={"filename": safe_name, "type": "file_digest"},
                    )
                    # Also update curated memory
                    twin.curated_memory.add_memory(f"File '{safe_name}': {summary[:500]}")
                    twin.curated_memory.refresh_snapshot()
                    logger.info("File digest complete: %s (%d char summary)", safe_name, len(summary))

                    # Emit event for frontend
                    if hasattr(twin, '_emit'):
                        twin._emit("file_digest", {
                            "filename": safe_name,
                            "summary_length": len(summary),
                        })
            except Exception as e:
                logger.warning("File digest failed for %s: %s", safe_name, e)

        # Fire and forget — don't block the upload response
        twin._bg_task(f"file-digest-{safe_name}", _file_digest())

    return {
        "filename": safe_name,
        "size": file_size,
        "type": ext,
        "content_preview": content_preview[:200],
        "stored": True,
    }


@app.get("/api/files/{filename}")
async def download_file(filename: str):
    """Download a file generated by the twin."""
    safe = filename.replace("/", "_").replace("\\", "_")
    file_path = OUTPUT_DIR / safe
    if not file_path.exists():
        # Also check uploads
        for f in UPLOAD_DIR.iterdir():
            if f.name.endswith(safe):
                file_path = f
                break
        else:
            return JSONResponse(status_code=404, content={"error": "File not found"})

    return FastAPIFileResponse(
        path=str(file_path),
        filename=safe,
        media_type="application/octet-stream",
    )


@app.get("/api/files")
async def list_files():
    """List available files for download."""
    files = []
    for f in sorted(OUTPUT_DIR.iterdir()):
        if f.is_file():
            files.append({"name": f.name, "size": f.stat().st_size, "type": "output"})
    for f in sorted(UPLOAD_DIR.iterdir()):
        if f.is_file():
            files.append({"name": f.name, "size": f.stat().st_size, "type": "upload"})
    return {"files": files}


@app.get("/api/stats/{session_id}")
async def get_stats(session_id: str):
    """Get current agent stats."""
    twin = twins.get(session_id)
    if not twin:
        return JSONResponse(status_code=404, content={"error": "Session not found"})

    return {
        "agent_id": twin.config.agent_id,
        "name": twin.config.name,
        "message_count": len(twin._messages),
        "events": events.get(session_id, []),
    }


# ── Serve frontend ──
FRONTEND_PATH = Path(__file__).parent / "web_ui.html"

@app.get("/")
async def serve_frontend():
    if FRONTEND_PATH.exists():
        return HTMLResponse(FRONTEND_PATH.read_text(encoding="utf-8"))
    return HTMLResponse("<h1>web_ui.html not found</h1><p>Place it next to web_demo.py</p>")


# ═══════════════════════════════════════════════════
# Startup
# ═══════════════════════════════════════════════════
def resolve_config(args):
    """Resolve config from args > env > defaults."""
    def env(key, default=""):
        return os.environ.get(key, default)

    provider = args.provider or env("TWIN_LLM_PROVIDER", "gemini")

    # API key
    api_key = args.api_key
    if not api_key:
        key_map = {"gemini": "GEMINI_API_KEY", "openai": "OPENAI_API_KEY", "anthropic": "ANTHROPIC_API_KEY"}
        api_key = env(key_map.get(provider, "GEMINI_API_KEY"))

    if not api_key:
        print("\n❌ No LLM API key found!")
        print(f"   Set GEMINI_API_KEY (or OPENAI_API_KEY) in .env or environment\n")
        sys.exit(1)

    net = args.network or env("NEXUS_NETWORK", "testnet")
    net_prefix = "MAINNET" if "mainnet" in net else "TESTNET"

    return {
        "provider": provider,
        "api_key": api_key,
        "model": args.model or env("TWIN_LLM_MODEL", ""),
        "name": env("TWIN_NAME", "Digital Twin"),
        "owner": env("TWIN_OWNER", ""),
        "private_key": env("NEXUS_PRIVATE_KEY", ""),
        "network": net,
        "rpc_url": env(f"NEXUS_{net_prefix}_RPC", ""),
        "agent_state_address": env(f"NEXUS_{net_prefix}_AGENT_STATE_ADDRESS", ""),
        "task_manager_address": env(f"NEXUS_{net_prefix}_TASK_MANAGER_ADDRESS", ""),
        "identity_registry": env(f"NEXUS_{net_prefix}_IDENTITY_REGISTRY", "") or env(f"NEXUS_{net_prefix}_IDENTITY_REGISTRY_ADDRESS", ""),
        "greenfield_bucket": env("NEXUS_GREENFIELD_BUCKET", "nexus-agent-state"),
        "tavily_key": env("TAVILY_API_KEY", ""),
        "jina_key": env("JINA_API_KEY", ""),
    }


def main():
    global CFG

    parser = argparse.ArgumentParser(description="Rune Protocol — Digital Twin Web Demo")
    parser.add_argument("--port", type=int, default=int(os.environ.get("PORT", "8000")))
    parser.add_argument("--provider", default="", choices=["gemini", "openai", "anthropic", ""])
    parser.add_argument("--model", default="")
    parser.add_argument("--api-key", default="")
    # Private key: env-only (NEXUS_PRIVATE_KEY) — never pass on CLI for security
    parser.add_argument("--network", default="", choices=["testnet", "mainnet", ""])
    args = parser.parse_args()

    CFG = resolve_config(args)

    # Load persisted users
    _load_users()

    chain_mode = "🔗 Chain" if CFG["private_key"] else "💻 Local"

    # Update WebAuthn origin to match actual port
    global ORIGIN, RP_ID
    ORIGIN = f"http://localhost:{args.port}"
    RP_ID = "localhost"

    passkey_status = "Passkey ✅" if HAS_WEBAUTHN else "Passkey ❌ (install py_webauthn)"

    print()
    print("╔═══════════════════════════════════════════════╗")
    print("║   Rune Protocol — Digital Twin Web Demo       ║")
    print("╠═══════════════════════════════════════════════╣")
    print(f"║   LLM:     {CFG['provider']} ({CFG['model'] or 'auto'})")
    print(f"║   Storage: {chain_mode} ({CFG['network']})")
    print(f"║   Auth:    {passkey_status}")
    print(f"║   URL:     http://localhost:{args.port}")
    print("╚═══════════════════════════════════════════════╝")
    print()

    # Console: INFO, File: DEBUG (captures everything for troubleshooting)
    log_file = PROJECT_ROOT / ".nexus_demo" / "demo.log"
    log_file.parent.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s [%(name)s] %(levelname)s: %(message)s",
        datefmt="%H:%M:%S",
    )
    file_handler = logging.FileHandler(str(log_file), mode="a", encoding="utf-8")
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(logging.Formatter(
        "%(asctime)s [%(name)s] %(levelname)s: %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    ))
    logging.getLogger().addHandler(file_handler)
    # Also capture SDK/Nexus internals
    for name in ["nexus", "nexus_core", "nexus_core.greenfield", "nexus_core.backend.chain"]:
        logging.getLogger(name).setLevel(logging.DEBUG)

    logger.info("Log file: %s", log_file)
    uvicorn.run(app, host="0.0.0.0", port=args.port, log_level="info")


if __name__ == "__main__":
    main()
